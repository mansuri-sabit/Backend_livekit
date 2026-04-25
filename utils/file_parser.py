"""Extract plain text from knowledge-base uploads: TXT, PDF, DOCX."""
import io
from pathlib import Path

from loguru import logger

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
MAX_FILE_SIZE_MB = 100


def extract_text_from_file(filename: str, content: bytes) -> str:
    """
    Extract text from an uploaded file.
    Raises ValueError for unsupported formats, oversized files, or parse failures.
    """
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported format: {ext}. Allowed: .txt, .pdf, .docx")
    if len(content) > MAX_FILE_SIZE_MB * 1024 * 1024:
        raise ValueError(f"File too large. Max {MAX_FILE_SIZE_MB} MB.")

    try:
        if ext == ".txt":
            return content.decode("utf-8", errors="replace").strip()

        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            parts = [page.extract_text() for page in reader.pages]
            return "\n\n".join(p for p in parts if p and isinstance(p, str)).strip()

        if ext == ".docx":
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts: list[str] = []

            # Extract paragraphs (body text, headings, list items)
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    parts.append(t)

            # Extract table cells — FAQ documents store Q&A pairs in tables
            # (Question column | Answer column). Without this, any DOCX whose
            # content lives in tables returns empty text and fails with
            # "No text extracted".
            for table in doc.tables:
                for row in table.rows:
                    # Join non-empty cells in each row with a tab separator
                    # so "Q: What is X?  A: X is Y" stays on one logical line
                    # rather than being split across two separate chunks.
                    row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # Deduplicate merged cells — python-docx repeats the text
                    # of merged cells for every cell they span.
                    seen: set[str] = set()
                    deduped: list[str] = []
                    for rp in row_parts:
                        if rp not in seen:
                            seen.add(rp)
                            deduped.append(rp)
                    if deduped:
                        parts.append("\t".join(deduped))

            return "\n\n".join(parts).strip()

    except ValueError:
        raise
    except Exception as exc:
        logger.error(f"[FileParser] Failed to parse {filename}: {exc}")
        raise ValueError(f"Could not extract text from {filename}: {exc}")

    return ""
